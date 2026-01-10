import os
import re
import json
import base64
import tempfile
import subprocess
import shutil
import hashlib
import streamlit as st
import xml.etree.ElementTree as ET
import requests
import time
import datetime
from io import BytesIO
import zipfile
from concurrent.futures import ThreadPoolExecutor
import uuid
from mistralai import Mistral
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image

# Cấu hình trang Streamlit - PHẢI LÀ LỆNH ĐẦU TIÊN
st.set_page_config(
    page_title="P_OCR PDF AI 2025",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS tùy chỉnh
st.markdown("""
<style>
    .main {
        padding: 1rem;
    }
    .stButton button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        font-weight: bold;
    }
    .card {
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 1rem;
        border: 1px solid #ddd;
    }
    .header-container {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 1rem;
    }
    .info-box {
        background-color: #e6f7ff;
        border-left: 4px solid #1890ff;
        padding: 1rem;
        margin-bottom: 1rem;
        border-radius: 4px;
    }
    .image-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(300px, 1fr));
        grid-gap: 16px;
    }
    .image-card {
        border: 1px solid #ddd;
        border-radius: 8px;
        padding: 8px;
        background-color: white;
    }
    .log-container {
        max-height: 300px;
        overflow-y: auto;
        background-color: #f9f9f9;
        padding: 10px;
        border-radius: 5px;
        border: 1px solid #ddd;
        font-family: 'Courier New', monospace;
        margin-bottom: 1rem;
    }
    .export-buttons {
        display: flex;
        flex-direction: column;
        gap: 10px;
        margin-top: 10px;
    }
    .export-buttons > div {
        margin-bottom: 10px;
    }
</style>
""", unsafe_allow_html=True)

# Định nghĩa hàm add_log sớm vì nó được sử dụng bởi nhiều hàm khác
def add_log(message):
    """Thêm log vào session_state"""
    timestamp = time.strftime("%H:%M:%S")
    if 'logs' not in st.session_state:
        st.session_state.logs = []
    st.session_state.logs.append(f"[{timestamp}] {message}")
    
    # Giữ log dưới một kích thước nhất định
    if len(st.session_state.logs) > 100:
        st.session_state.logs = st.session_state.logs[-100:]

# Hàm hiển thị logs trong container
def display_logs():
    """Hiển thị logs trong container"""
    if 'logs' in st.session_state and st.session_state.logs:
        logs_str = "\n".join(st.session_state.logs)
        st.markdown(f"""
        <div class="log-container">
        {logs_str.replace("\n", "<br>")}
        </div>
        """, unsafe_allow_html=True)
    else:
        st.info("Chưa có nhật ký hoạt động")

def load_rsa_private_key_from_xml(xml_str):
    """
    Tải khóa RSA riêng tư từ định dạng XML 
    """
    try:
        import Crypto
        from Crypto.PublicKey import RSA
        
        root = ET.fromstring(xml_str)
        def get_int(tag):
            text = root.find(tag).text
            return int.from_bytes(base64.b64decode(text), 'big')
        n = get_int('Modulus')
        e = get_int('Exponent')
        d = get_int('D')
        p = get_int('P')
        q = get_int('Q')
        key = RSA.construct((n, e, d, p, q))
        return key
    except ImportError:
        st.error("Thư viện Crypto không được cài đặt. Vui lòng sử dụng API key thay thế.")
        return None
    except Exception as e:
        st.error(f"Lỗi khi tải khóa RSA: {str(e)}")
        return None

def decrypt_api_key(encrypted_key_base64, rsa_private_key):
    """
    Giải mã API key đã được mã hóa sử dụng khóa RSA riêng tư
    """
    try:
        from Crypto.Cipher import PKCS1_v1_5
        
        cipher = PKCS1_v1_5.new(rsa_private_key)
        
        # Giải mã Base64 để có dữ liệu nhị phân
        encrypted_data = base64.b64decode(encrypted_key_base64)
        
        # Giải mã RSA
        decrypted = cipher.decrypt(encrypted_data, None)
        
        if not decrypted:
            raise ValueError("Giải mã thất bại")
            
        # Chuyển kết quả giải mã thành chuỗi
        return decrypted.decode('utf-8')
        
    except Exception as e:
        raise ValueError(f"Lỗi giải mã API key: {str(e)}")

def get_mineru_token():
    """
    Lấy API key đã được mã hóa từ Github và giải mã bằng khóa RSA riêng tư
    """
    try:
        # Khóa RSA riêng tư để giải mã API key
        PRIVATE_KEY_XML = """<RSAKeyValue>
<Modulus>pWVItQwZ7NCPcBhSL4rqJrwh4OQquiPVtqTe4cqxO7o+UjYNzDPfLkfKAvR8k9ED4lq2TU11zEj8p2QZAM7obUlK4/HVexzfZd0qsXlCy5iaWoTQLXbVdzjvkC4mkO5TaX3Mpg/+p4oZjk1iS68tQFmju5cT19dcsPh554ICk8U=</Modulus>
<Exponent>AQAB</Exponent>
<P>0ZWwsKa9Vw9BJAsRaW4eV60i6Z+R6z9LNSgjNn4pYH2meZtGUbmJVowRv7EM5sytouB5EMru7sQbRHEQ7nrwSw==</P>
<Q>ygZQWNkUgfHhHBataXvYLxWgPB5UZTWogN8Mb33LT4rq7I5P1GX3oWtYF2AdmChX8Lq3Ms/A/jBhqYomhYOiLw==</Q>
<DP>qS9VOsTfA3Bk/VuR6rHh/JTfIgiWGnk1lOuZwVuGu0WzJWebFE3Z9+uKSFv8NjPz1w+tq0imKEhWWqGLMXg8kQ==</DP>
<DQ>UCtXQRrMB5EL6tCY+k4aCP1E+/ZxOUSk3Jcm4SuDPcp71WnYBgp8zULCz2vl8pa35yDBSFmnVXevmc7n4H3PIw==</DQ>
<InverseQ>Qm9RjBhxANWyIb8I28vjGz+Yb9CnunWxpHWbfRo1vF+Z38WB7dDgLsulAXMGrUPQTeG6K+ot5moeZ9ZcAc1Hzw==</InverseQ>
<D>F9lU9JY8HsOsCzPWlfhn7xHtqKn95z1HkcCQSuqZR82BMwWMU8efBONhI6/xTrcy4i7GXrsuozhbBiAO4ujy5qPytdFemLuqjwFTyvllkcOy3Kbe0deczxnPPCwmSMVKsYInByJoBP3JYoyVAj4bvY3UqZJtw+2u/OIOhoBe33k=</D>
</RSAKeyValue>"""
        
        # Tải khóa RSA riêng tư từ XML
        rsa_private_key = load_rsa_private_key_from_xml(PRIVATE_KEY_XML)
        if not rsa_private_key:
            return None
        
        # URL đến file chứa API key đã mã hóa trên Github
        github_url = "https://raw.githubusercontent.com/thayphuctoan/pconvert/refs/heads/main/ocr-pdf-hvt"
        
        # Lấy nội dung từ Github
        response = requests.get(github_url, timeout=10)
        response.raise_for_status()  # Đảm bảo yêu cầu thành công
        
        # Lấy các dòng không trống từ phản hồi
        encrypted_keys = [line.strip() for line in response.text.splitlines() if line.strip()]
        
        if not encrypted_keys:
            raise ValueError("Không tìm thấy API key đã mã hóa")
        
        # Chỉ lấy token đầu tiên từ danh sách và giải mã
        token = decrypt_api_key(encrypted_keys[0], rsa_private_key)
        
        # Kiểm tra nếu token rỗng
        if not token:
            raise ValueError("API key giải mã rỗng")
            
        return token
    
    except requests.RequestException as e:
        st.error(f"Lỗi kết nối đến Github: {str(e)}")
    except ValueError as e:
        st.error(f"Lỗi dữ liệu: {str(e)}")
    except Exception as e:
        st.error(f"Lỗi lấy API key: {str(e)}")
    
    return None

def count_pdf_pages(pdf_file):
    """Đếm số trang trong file PDF"""
    try:
        pdf = PdfReader(pdf_file)
        return len(pdf.pages)
    except Exception as e:
        st.error(f"Lỗi khi đọc file PDF: {str(e)}")
        return -1

def process_ocr(pdf_file):
    """Xử lý OCR cho file PDF đã upload"""
    try:
        # Lưu file tạm thời để xử lý
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(pdf_file.getvalue())
            tmp_file_path = tmp_file.name
            
        add_log(f"Bắt đầu xử lý OCR cho file: {pdf_file.name}")
        
        # Tự động lấy Mistral API key
        add_log("Đang lấy Server API key tự động...")
        mistral_api_key = get_mineru_token()
        
        if not mistral_api_key:
            add_log("Không thể lấy Server API key tự động")
            st.error("Không thể lấy Server API key tự động. Vui lòng thử lại sau.")
            return {"error": "Không thể lấy Server API key"}
        
        add_log("Đã lấy server API key thành công")
        
        # Khởi tạo client Mistral
        client = Mistral(api_key=mistral_api_key)
        
        add_log("Đang upload file PDF...")
        uploaded_pdf = client.files.upload(
            file={
                "file_name": pdf_file.name,
                "content": open(tmp_file_path, "rb"),
            },
            purpose="ocr"
        )
        
        add_log(f"File đã được upload với ID: {uploaded_pdf.id}")
        
        add_log("Đang lấy thông tin file...")
        file_info = client.files.retrieve(file_id=uploaded_pdf.id)
        add_log(f"Thông tin file: {file_info}")
        
        add_log("Đang lấy signed URL...")
        signed_url = client.files.get_signed_url(file_id=uploaded_pdf.id)
        
        add_log("Đang xử lý OCR...")
        ocr_response = client.ocr.process(
            model="mistral-ocr-latest",
            document={
                "type": "document_url",
                "document_url": signed_url.url,
            },
            include_image_base64=True
        )
        
        # Lưu toàn bộ phản hồi để gỡ lỗi
        raw_response_str = str(ocr_response)
        add_log(f"Kích thước phản hồi: {len(raw_response_str)} ký tự")
        
        result_data = {
            "text": "",
            "images": {},
            "raw_response": raw_response_str
        }
        
        # Trích xuất văn bản và hình ảnh từ phản hồi API
        text_extracted = False  # Thêm biến cờ để kiểm soát việc trích xuất
        
        # Phương pháp 1: Trích xuất từ model_dump
        if hasattr(ocr_response, 'model_dump') and not text_extracted:
            try:
                response_dict = ocr_response.model_dump()
                extract_from_dict(response_dict, result_data)
                # Kiểm tra xem đã trích xuất thành công hay chưa
                if result_data["text"]:
                    text_extracted = True
                    add_log("Đã trích xuất văn bản qua model_dump")
            except Exception as e:
                add_log(f"Lỗi khi phân tích model_dump: {str(e)}")
        
        # Phương pháp 2: Trích xuất trực tiếp từ thuộc tính pages
        if hasattr(ocr_response, 'pages') and not text_extracted:
            try:
                pages = ocr_response.pages
                for page in pages:
                    # Trích xuất văn bản
                    if hasattr(page, 'markdown') and page.markdown:
                        result_data["text"] += page.markdown + "\n\n"
                    elif hasattr(page, 'text') and page.text:
                        result_data["text"] += page.text + "\n\n"
                    
                    # Trích xuất hình ảnh (điều này vẫn được thực hiện bất kể text_extracted)
                    if hasattr(page, 'images') and page.images:
                        for img in page.images:
                            if hasattr(img, 'id') and hasattr(img, 'image_base64'):
                                img_id = img.id
                                img_base64 = img.image_base64
                                result_data["images"][img_id] = img_base64
                                add_log(f"Tìm thấy hình ảnh {img_id} (kích thước: {len(img_base64)} ký tự)")
                
                if result_data["text"]:
                    text_extracted = True
                    add_log("Đã trích xuất văn bản trực tiếp từ thuộc tính pages")
            except Exception as e:
                add_log(f"Lỗi khi trích xuất từ pages: {str(e)}")
        
        # Phương pháp 3: Dùng regex nếu hai phương pháp trên thất bại
        if not result_data["text"] or not result_data["images"]:
            add_log("Sử dụng regex để trích xuất dữ liệu")
            extract_with_regex(raw_response_str, result_data)
        
        # In ra thông tin hình ảnh đã tìm thấy
        if result_data["images"]:
            for img_id, base64_data in result_data["images"].items():
                add_log(f"Hình ảnh {img_id}: {len(base64_data)} ký tự")
        else:
            add_log("Không tìm thấy hình ảnh trong phản hồi API")
        
        add_log("Đang xử lý định dạng kết quả...")
        
        # Làm sạch văn bản trước khi trả về
        text = result_data["text"]
        
        # Làm sạch văn bản
        cleaned_text = text
        
        # Loại bỏ các ký tự không cần thiết
        cleaned_text = re.sub(r'OCRPageObject\(.*?\)', '', cleaned_text)
        cleaned_text = re.sub(r'OCRPageDimensions\(.*?\)', '', cleaned_text)
        cleaned_text = re.sub(r'images=\[\]', '', cleaned_text)
        cleaned_text = re.sub(r'index=\d+', '', cleaned_text)
        
        # Tiền xử lý để thêm khoảng cách trước các mục quan trọng
        cleaned_text = re.sub(r'(Câu\s+\d+\.?[:]?)', r'\n\n\1', cleaned_text)
        cleaned_text = re.sub(r'(Bài\s+\d+\.?[:]?)', r'\n\n\1', cleaned_text)
        cleaned_text = re.sub(r'([A-D]\.)', r'\n\1', cleaned_text)
        
        # Làm sạch các tham chiếu hình ảnh lặp
        # Xử lý tham chiếu dạng ![[HÌNH: img-X.jpeg]]([HÌNH: img-X.jpeg])
        cleaned_text = re.sub(r'!\[\[HÌNH: (img-\d+\.jpeg)\]\]\(\[HÌNH: \1\]\)', r'[HÌNH: \1]', cleaned_text)
        
        # Xử lý tham chiếu dạng ![img-X.jpeg](img-X.jpeg)
        cleaned_text = re.sub(r'!\[(img-\d+\.jpeg)\]\(\1\)', r'[HÌNH: \1]', cleaned_text)
        
        # Xử lý các tham chiếu hình ảnh khác có thể lặp
        for img_id in result_data["images"].keys():
            # Chuẩn hóa tất cả các tham chiếu thành [HÌNH: img-X.jpeg]
            
            # Tìm và thay thế dạng ![]()
            pattern = r'!\[.*?\]\(.*?' + re.escape(img_id) + r'.*?\)'
            cleaned_text = re.sub(pattern, f'[HÌNH: {img_id}]', cleaned_text)
            
            # Tìm và thay thế markdown inline image khác
            pattern = r'!{1,2}\[' + re.escape(img_id) + r'\]'
            cleaned_text = re.sub(pattern, f'[HÌNH: {img_id}]', cleaned_text)
            
            # Thay thế các đề cập trực tiếp đến ID hình ảnh (chỉ khi nó đứng riêng)
            pattern = r'(?<![a-zA-Z0-9\-\.])' + re.escape(img_id) + r'(?![a-zA-Z0-9\-\.])'
            cleaned_text = re.sub(pattern, f'[HÌNH: {img_id}]', cleaned_text)
            
            # Loại bỏ các dòng chỉ chứa ID hình ảnh
            pattern = r'^\s*' + re.escape(img_id) + r'\s*$'
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.MULTILINE)
        
        # Loại bỏ các dòng có chứa cụm "HÌNH:" hoặc "img-" mà không nằm trong cặp ngoặc vuông
        cleaned_text = re.sub(r'^\s*HÌNH:\s*.*$', '', cleaned_text, flags=re.MULTILINE)
        cleaned_text = re.sub(r'^\s*img-\d+\.jpeg\s*$', '', cleaned_text, flags=re.MULTILINE)
        
        # Loại bỏ các dấu ngoặc vuông đơn lẻ
        cleaned_text = re.sub(r'^\s*\]\s*$', '', cleaned_text, flags=re.MULTILINE)
        
        # Loại bỏ nhiều dòng trống liên tiếp
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        # Lưu văn bản đã làm sạch
        result_data["cleaned_text"] = cleaned_text
        
        # Xóa file tạm
        try:
            os.unlink(tmp_file_path)
        except:
            pass
            
        return result_data
        
    except Exception as e:
        add_log(f"Lỗi trong quá trình xử lý OCR: {str(e)}")
        import traceback
        add_log(traceback.format_exc())
        return {"error": str(e)}

def extract_from_dict(data_dict, result_data):
    """Trích xuất dữ liệu từ một dictionary"""
    if 'pages' in data_dict:
        for page in data_dict['pages']:
            # Trích xuất văn bản
            if 'markdown' in page:
                result_data["text"] += page['markdown'] + "\n\n"
            elif 'text' in page:
                result_data["text"] += page['text'] + "\n\n"
            
            # Trích xuất hình ảnh
            if 'images' in page and page['images']:
                for img in page['images']:
                    if 'id' in img and 'image_base64' in img:
                        img_id = img['id']
                        img_base64 = img['image_base64']
                        result_data["images"][img_id] = img_base64
                        add_log(f"Tìm thấy hình ảnh {img_id} (dict)")

def extract_with_regex(text, result_data):
    """Trích xuất dữ liệu bằng regex"""
    # Tìm tất cả các đoạn markdown
    markdown_blocks = re.findall(r'"markdown":\s*"(.*?)"(?=,|\})', text, re.DOTALL)
    for block in markdown_blocks:
        # Thay thế các ký tự escape
        block = block.replace('\\n', '\n').replace('\\"', '"')
        result_data["text"] += block + "\n\n"
    
    # Tìm tất cả các cặp id và image_base64
    image_matches = re.findall(r'"id":\s*"(img-\d+\.jpeg)",.*?"image_base64":\s*"([^"]+)"', text, re.DOTALL)
    for img_id, img_base64 in image_matches:
        result_data["images"][img_id] = img_base64
        add_log(f"Tìm thấy hình ảnh {img_id} (regex)")
    
    # Nếu không tìm thấy theo cặp, thử tìm riêng
    if not result_data["images"]:
        img_ids = re.findall(r'"id":\s*"(img-\d+\.jpeg)"', text)
        base64_data = re.findall(r'"image_base64":\s*"([^"]+)"', text)
        
        if len(img_ids) == len(base64_data):
            for i in range(len(img_ids)):
                result_data["images"][img_ids[i]] = base64_data[i]
                add_log(f"Tìm thấy hình ảnh {img_ids[i]} (regex riêng)")

def process_formulas(text):
    """Xử lý và chuẩn hóa công thức toán học trong văn bản"""
    def process_math_content(match):
        content = match.group(1)
        content = content.replace('π', '\\pi')
        content = re.sub(r'√(\d+)', r'\\sqrt{\1}', content)
        content = re.sub(r'√\{([^}]+)\}', r'\\sqrt{\1}', content)
        content = content.replace('≠', '\\neq')
        content = content.replace('*', '')
        return f'${content}$'

    text = re.sub(r'\$(.+?)\$', process_math_content, text, flags=re.DOTALL)
    return text

def check_pandoc_installed():
    """Kiểm tra xem pandoc đã được cài đặt chưa"""
    try:
        result = subprocess.run(["pandoc", "--version"], 
                               stdout=subprocess.PIPE, 
                               stderr=subprocess.PIPE, 
                               text=True, 
                               check=False)
        installed = result.returncode == 0
        if installed:
            add_log("Pandoc đã được cài đặt: " + result.stdout.split('\n')[0])
        else:
            add_log("Pandoc không được cài đặt hoặc không khả dụng")
        return installed
    except FileNotFoundError:
        add_log("Pandoc không được cài đặt (FileNotFoundError)")
        return False
    except Exception as e:
        add_log(f"Lỗi khi kiểm tra Pandoc: {str(e)}")
        return False

def generate_word_document(text, images):
    """Tạo file Word từ văn bản và hình ảnh với độ tin cậy cao"""
    try:
        # Tạo thư mục tạm thời
        temp_dir = tempfile.mkdtemp()
        add_log(f"Đã tạo thư mục tạm thời: {temp_dir}")
        
        # Chuẩn bị nội dung markdown từ text
        markdown_text = text
        
        # Xử lý công thức toán học (nếu có)
        markdown_text = process_formulas(markdown_text)
        
        # Tạo file CSS tùy chỉnh để cải thiện định dạng
        css_file = os.path.join(temp_dir, "custom.css")
        with open(css_file, "w", encoding="utf-8") as f:
            f.write("""
            body {
                font-family: 'Times New Roman', serif;
                font-size: 12pt;
                line-height: 1.5;
                margin: 2.5cm;
            }
            h1, h2, h3, h4, h5, h6 {
                font-family: 'Arial', sans-serif;
                color: #2c3e50;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
            }
            h1 { font-size: 18pt; }
            h2 { font-size: 16pt; }
            h3 { font-size: 14pt; }
            p { margin-bottom: 0.8em; text-align: justify; }
            img { 
                display: block; 
                margin: 1em auto; 
                max-width: 95%;
                page-break-inside: avoid;
            }
            .caption { 
                text-align: center;
                font-style: italic;
                margin-top: 0.3em;
                color: #666;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 1em 0;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8px;
            }
            th {
                background-color: #f2f2f2;
                font-weight: bold;
            }
            """)
        add_log("Đã tạo CSS tùy chỉnh cho định dạng Word đẹp hơn")
        
        # Lưu và xử lý hình ảnh với độ tin cậy cao hơn
        image_paths = {}
        for img_id, base64_data in images.items():
            try:
                # Tạo đường dẫn đầy đủ cho file hình ảnh
                img_path = os.path.join(temp_dir, img_id)
                
                # Xử lý dữ liệu base64
                if "," in base64_data:
                    # Nếu dữ liệu có định dạng data URI (data:image/jpeg;base64,...)
                    base64_data = base64_data.split(",", 1)[1]
                
                # Giải mã base64 và lưu vào file
                img_bytes = base64.b64decode(base64_data)
                with open(img_path, "wb") as f:
                    f.write(img_bytes)
                
                # Kiểm tra xem file có phải là hình ảnh hợp lệ không
                try:
                    from PIL import Image
                    img = Image.open(img_path)
                    # Lưu lại với định dạng rõ ràng để tránh vấn đề với Pandoc
                    if img.mode == 'RGBA':
                        # Chuyển đổi hình ảnh PNG RGBA thành RGB để tránh vấn đề tương thích
                        rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                        rgb_img.paste(img, mask=img.split()[3])  # Sử dụng kênh alpha làm mask
                        rgb_img.save(img_path, 'JPEG', quality=90)
                    else:
                        img.save(img_path, quality=90)
                    img.close()
                    add_log(f"Đã xác minh và tối ưu hình ảnh {img_id}")
                except Exception as img_err:
                    add_log(f"Cảnh báo khi xử lý hình ảnh {img_id}: {str(img_err)}, tiếp tục với dữ liệu gốc")
                
                image_paths[img_id] = img_path
                add_log(f"Đã lưu hình ảnh {img_id} vào {img_path}")
                
            except Exception as e:
                add_log(f"Lỗi khi lưu hình ảnh {img_id}: {str(e)}")
        
        # Chèn hình ảnh vào vị trí đã đánh dấu với định dạng phong phú hơn
        for img_id, img_path in image_paths.items():
            rel_path = os.path.basename(img_path)
            placeholder = f"[HÌNH: {img_id}]"
            
            # Format tham chiếu hình ảnh đúng cú pháp Markdown với caption
            image_ref = f"""
![{img_id}]({rel_path})

<div class="caption">Hình {img_id.replace('img-', '').replace('.jpeg', '')}</div>
"""
            
            # Tìm và thay thế tất cả các placeholder
            if placeholder in markdown_text:
                # Đảm bảo hình ảnh hiển thị riêng dòng với khoảng cách phù hợp
                markdown_text = markdown_text.replace(placeholder, image_ref)
            else:
                # Nếu không tìm thấy placeholder, thêm vào cuối
                add_log(f"Không tìm thấy placeholder cho hình ảnh {img_id}, thêm vào cuối tài liệu")
                markdown_text += f"\n\n{image_ref}\n\n"
        
        # Lưu markdown vào file tạm thời
        md_file = os.path.join(temp_dir, "export.md")
        with open(md_file, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        
        add_log(f"Đã lưu nội dung markdown vào {md_file}")
        
        # Kiểm tra xem pandoc đã được cài đặt chưa
        pandoc_installed = check_pandoc_installed()
        
        if not pandoc_installed:
            add_log("Pandoc không được cài đặt. Đang thử dùng python-docx thay thế...")
            
            # Thử sử dụng python-docx nếu có
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Tạo file docx
                doc = Document()
                
                # Thiết lập style
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)
                
                # Thêm tiêu đề
                doc.add_heading('Kết quả OCR', 0)
                
                # Thêm nội dung
                paragraphs = markdown_text.split('\n\n')
                for para_text in paragraphs:
                    if not para_text.strip():
                        continue
                        
                    # Kiểm tra xem đoạn văn có phải là tham chiếu hình ảnh không
                    img_match = re.search(r'!\[(.*?)\]\((.*?)\)', para_text)
                    if img_match:
                        img_id = img_match.group(1)
                        img_path = img_match.group(2)
                        # Tìm đường dẫn đầy đủ của hình ảnh
                        for id, path in image_paths.items():
                            if os.path.basename(path) == os.path.basename(img_path):
                                # Thêm hình ảnh với kích thước phù hợp
                                picture = doc.add_picture(path, width=Inches(6))
                                # Thêm caption
                                caption = doc.add_paragraph(f"Hình {img_id.replace('img-', '').replace('.jpeg', '')}")
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption.style = 'Caption'
                                break
                    else:
                        # Thêm đoạn văn bản thông thường
                        if para_text.startswith('# '):
                            # Heading 1
                            doc.add_heading(para_text[2:], 1)
                        elif para_text.startswith('## '):
                            # Heading 2
                            doc.add_heading(para_text[3:], 2)
                        elif para_text.startswith('### '):
                            # Heading 3
                            doc.add_heading(para_text[4:], 3)
                        elif para_text.startswith('$') and para_text.endswith('$'):
                            # Công thức toán học inline
                            p = doc.add_paragraph()
                            p.add_run(para_text)
                            p.italic = True
                        elif para_text.startswith('$$') and para_text.endswith('$$'):
                            # Công thức toán học block
                            p = doc.add_paragraph()
                            p.add_run(para_text[2:-2])
                            p.italic = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # Đoạn văn thông thường
                            doc.add_paragraph(para_text)
                
                # Lưu file docx
                docx_file = os.path.join(temp_dir, "output.docx")
                doc.save(docx_file)
                
                # Đọc file docx đã tạo
                with open(docx_file, "rb") as f:
                    docx_data = f.read()
                
                add_log("Đã tạo file Word thành công bằng python-docx")
                
                # Dọn dẹp
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
                
                return {
                    "docx_data": docx_data,
                    "is_docx": True,
                    "filename": "ocr_result.docx"
                }
                
            except ImportError:
                add_log("python-docx không được cài đặt. Đang xuất file Markdown thay thế.")
                # Tạo ZIP chứa markdown và hình ảnh
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    # Thêm file markdown
                    zip_file.write(md_file, os.path.basename(md_file))
                    
                    # Thêm tất cả hình ảnh
                    for img_id, img_path in image_paths.items():
                        zip_file.write(img_path, os.path.basename(img_path))
                    
                    # Thêm file HTML để dễ xem
                    html_content = f"""<!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>Kết quả OCR</title>
                        <style>
                            body {{ font-family: 'Times New Roman', serif; line-height: 1.6; margin: 0 auto; max-width: 800px; padding: 20px; }}
                            img {{ max-width: 100%; height: auto; display: block; margin: 20px auto; }}
                            h1, h2 {{ color: #2c3e50; }}
                            pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                            .caption {{ text-align: center; font-style: italic; color: #777; margin-top: 5px; }}
                        </style>
                    </head>
                    <body>
                        <h1>Kết quả OCR</h1>
                        {markdown_text.replace('[HÌNH: ', '<img src="').replace(']', '"><div class="caption">Hình </div>')}
                    </body>
                    </html>"""
                    
                    zip_file.writestr("preview.html", html_content)
                
                # Reset buffer để đọc từ đầu
                zip_buffer.seek(0)
                
                # Dọn dẹp
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
                
                return {
                    "zip_data": zip_buffer.getvalue(),
                    "is_docx": False,
                    "filename": "ocr_result.zip"
                }
            except Exception as e:
                add_log(f"Lỗi khi sử dụng python-docx: {str(e)}")
                # Quay lại cách xuất zip nếu có lỗi
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    # Thêm file markdown
                    zip_file.write(md_file, os.path.basename(md_file))
                    
                    # Thêm tất cả hình ảnh
                    for img_id, img_path in image_paths.items():
                        zip_file.write(img_path, os.path.basename(img_path))
                
                zip_buffer.seek(0)
                
                # Dọn dẹp
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
                
                return {
                    "zip_data": zip_buffer.getvalue(),
                    "is_docx": False,
                    "filename": "ocr_result.zip"
                }
        
        # Nếu Pandoc được cài đặt, tiếp tục sử dụng
        add_log("Đang chạy Pandoc để chuyển đổi sang Word...")
        
        docx_file = os.path.join(temp_dir, "output.docx")
        
        # Sử dụng tham số đặc biệt để cải thiện chất lượng đầu ra
        pandoc_command = [
            "pandoc", 
            md_file, 
            "-o", docx_file, 
            "--resource-path", temp_dir,
            "--css", css_file,
            "--reference-doc=reference.docx" if os.path.exists("reference.docx") else "",
            "--mathjax",         # Hỗ trợ công thức toán học
            "--toc",             # Tạo mục lục tự động
            "--toc-depth=3",     # Độ sâu mục lục
            "--standalone"       # Tạo tài liệu độc lập
        ]
        
        # Loại bỏ các tham số trống
        pandoc_command = [cmd for cmd in pandoc_command if cmd]
        
        add_log(f"Lệnh Pandoc: {' '.join(pandoc_command)}")
        
        # Chạy lệnh Pandoc
        result = subprocess.run(
            pandoc_command,
            check=True, 
            capture_output=True, 
            text=True
        )
        
        # Đọc file Word đã tạo
        with open(docx_file, "rb") as f:
            docx_data = f.read()
        
        add_log("Đã tạo file Word thành công bằng Pandoc")
        
        # Dọn dẹp
        try:
            shutil.rmtree(temp_dir)
        except:
            pass
        
        return {
            "docx_data": docx_data,
            "is_docx": True,
            "filename": "ocr_result.docx"
        }
        
    except Exception as e:
        add_log(f"Lỗi khi tạo file Word: {str(e)}")
        import traceback
        add_log(traceback.format_exc())
        
        # Trả về lỗi
        return {
            "error": str(e),
            "is_docx": False
        }

# ----- Ứng dụng Streamlit chính -----

# Tạo tiêu đề và giới thiệu
st.title("HỆ THỐNG GIÁO DỤC HOÀNG TÂY BẮC")
st.markdown("<p> PHẦN MỀM CHUYỂN PDF SANG WORD </p>", unsafe_allow_html=True)

# Sidebar với phần cài đặt và thông tin
with st.sidebar:
    st.header("HƯỚNG DẪN SỬ DỤNG")
    st.info("""
    1. Upload file PDF
    2. Nhấn nút xử lý OCR
    3. Xem và tải kết quả
    """)
    
    st.subheader("Thông tin liên hệ")
    st.markdown("### ThS. Hà Huy Hoàng")  
    st.markdown("#### Liên hệ: 097.125.8386")
# Main container for the application
main_container = st.container()

with main_container:
    # File upload section
    st.subheader("Upload File PDF")
    
    uploaded_file = st.file_uploader("Chọn file PDF để xử lý OCR", type=["pdf"])
    
    if uploaded_file:
        # Hiển thị thông tin file
        file_details = {"Tên file": uploaded_file.name, "Loại file": uploaded_file.type, "Kích thước": f"{uploaded_file.size/1024:.1f} KB"}
        
        # Đếm số trang
        page_count = count_pdf_pages(uploaded_file)
        if page_count > 0:
            file_details["Số trang"] = page_count
            
            # Kiểm tra giới hạn số trang
            if page_count > 500:
                st.error(f"File PDF có {page_count} trang, vượt quá giới hạn 100 trang. Vui lòng chọn file nhỏ hơn hoặc chia nhỏ file này.")
        
        # Hiển thị thông tin file dưới dạng expander
        with st.expander("Thông tin file", expanded=True):
            for key, value in file_details.items():
                st.write(f"**{key}:** {value}")

    # Processing section
    st.subheader("Xử lý OCR")

    # Nút xử lý OCR
    process_button = st.button(
        "Xử lý OCR",
        disabled=not uploaded_file or (page_count > 500 if page_count > 0 else False),
        key="process_ocr_button"
    )
    
    # Xử lý OCR khi nhấn nút
    if process_button:
        if page_count > 500:
            st.error(f"File PDF có {page_count} trang, vượt quá giới hạn 100 trang. Vui lòng chọn file nhỏ hơn.")
        else:
            # Đặt lại logs
            st.session_state.logs = []
            
            # Đặt lại kết quả trước đó
            if "result_data" in st.session_state:
                del st.session_state.result_data
            if "docx_result" in st.session_state:
                del st.session_state.docx_result
            
            add_log(f"Bắt đầu xử lý OCR cho file: {uploaded_file.name}")
            
            with st.spinner("Đang xử lý OCR... (quá trình này có thể mất vài phút)"):
                # Xử lý OCR
                result = process_ocr(uploaded_file)
                
                # Lưu kết quả vào session_state
                st.session_state.result_data = result
                
                if "error" in result:
                    st.error(f"Lỗi: {result['error']}")
                else:
                    st.success("Xử lý OCR hoàn tất thành công!")
    
    # Hiển thị logs
    st.subheader("Nhật ký hoạt động")
    display_logs()
    
    # Hiển thị kết quả OCR
    if "result_data" in st.session_state and st.session_state.result_data and "error" not in st.session_state.result_data:
        st.subheader("Kết quả OCR")
        
        result_data = st.session_state.result_data
        
        # Tạo tabs để hiển thị kết quả văn bản và hình ảnh
        result_tabs = st.tabs(["Văn bản", "Hình ảnh"])
        
        # Tab văn bản
        with result_tabs[0]:
            # Hiển thị văn bản
            st.text_area("Kết quả OCR:", result_data["cleaned_text"], height=500)
            
            # Các nút tải xuống ngang hàng
            col1, col2 = st.columns(2)
            
            with col1:
                # Tải xuống văn bản
                st.download_button(
                    label="Tải văn bản (TXT)",
                    data=result_data["cleaned_text"],
                    file_name="ocr_text.txt",
                    mime="text/plain"
                )
            
            with col2:
                # Nút tạo file Word
                if "docx_result" not in st.session_state:
                    if st.button("Tạo file Word"):
                        with st.spinner("Đang tạo file Word..."):
                            # Tạo file Word
                            st.session_state.docx_result = generate_word_document(
                                result_data["cleaned_text"], 
                                result_data["images"]
                            )
                            st.rerun()
                else:
                    # Hiển thị nút tải xuống Word
                    if st.session_state.docx_result.get("is_docx", False):
                        st.download_button(
                            label="Tải file Word",
                            data=st.session_state.docx_result["docx_data"],
                            file_name=st.session_state.docx_result["filename"],
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    elif "zip_data" in st.session_state.docx_result:
                        st.download_button(
                            label="Tải file Markdown + hình ảnh (ZIP)",
                            data=st.session_state.docx_result["zip_data"],
                            file_name=st.session_state.docx_result["filename"],
                            mime="application/zip"
                        )
                    elif "error" in st.session_state.docx_result:
                        st.error(f"Lỗi khi tạo file Word: {st.session_state.docx_result['error']}")
        
        # Tab hình ảnh
        with result_tabs[1]:
            images = result_data.get("images", {})
            
            if images:
                st.write(f"Đã tìm thấy {len(images)} hình ảnh:")
                
                # Hiển thị hình ảnh trong lưới
                st.markdown('<div class="image-grid">', unsafe_allow_html=True)
                
                for i, (img_id, img_data) in enumerate(images.items()):
                    col = i % 3
                    
                    # Xử lý dữ liệu base64
                    if "," in img_data:
                        img_data = img_data.split(",", 1)[1]
                    
                    # Hiển thị hình ảnh
                    st.markdown(f"""
                    <div class="image-card">
                        <h4>{img_id}</h4>
                        <img src="data:image/jpeg;base64,{img_data}" alt="{img_id}" style="width:100%">
                    </div>
                    """, unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Tạo ZIP chứa tất cả hình ảnh
                img_zip_buffer = BytesIO()
                with zipfile.ZipFile(img_zip_buffer, 'w') as img_zip:
                    for img_id, img_data in images.items():
                        # Xử lý dữ liệu base64
                        if "," in img_data:
                            img_data = img_data.split(",", 1)[1]
                        
                        # Thêm hình ảnh vào ZIP
                        img_zip.writestr(img_id, base64.b64decode(img_data))
                
                img_zip_buffer.seek(0)
                
                # Nút tải xuống tất cả hình ảnh dưới dạng ZIP
                st.download_button(
                    label="Tải tất cả hình ảnh (ZIP)",
                    data=img_zip_buffer,
                    file_name="ocr_images.zip",
                    mime="application/zip"
                )
            else:
                st.info("Không tìm thấy hình ảnh trong kết quả OCR.")
